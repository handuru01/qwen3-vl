"""
개발보고서 Word 문서 생성 스크립트
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="맑은 고딕", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 한글 폰트 설정
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(31, 73, 125))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        set_font(run, size=13, bold=True, color=(31, 78, 121))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        set_font(run, size=11, bold=True, color=(68, 68, 68))
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def add_paragraph(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=10)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # 회색 배경 효과를 위해 음영 설정
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    rPr.append(shd)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더 행
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        # 헤더 배경색 (진한 파란색)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "EAF0FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

    doc.add_paragraph()
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


# ── 문서 생성 ──────────────────────────────────────────────────────────────────
doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── 표지 ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nQwen3-VL 멀티모달 AI 채팅\n웹서비스 개발보고서")
set_font(run, size=24, bold=True, color=(31, 73, 125))
p.paragraph_format.space_after = Pt(20)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("로컬 AI로 이미지 + 텍스트를 동시에 대화하는 웹서비스 구축")
set_font(run2, size=12, color=(89, 89, 89))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("\n\n개발자: 김상민\n작성일: 2026년 5월 11일\n")
set_font(run3, size=11, color=(89, 89, 89))

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("서비스 URL: https://atpia.com/qwen3-vl/index.html\n"
                   "소스코드: https://github.com/handuru01/qwen3-vl")
set_font(run4, size=10, color=(31, 73, 125))

doc.add_page_break()

# ── 1. 개요 ───────────────────────────────────────────────────────────────────
add_heading(doc, "1. 개요")
add_divider(doc)
add_paragraph(doc,
    "로컬 PC에 설치된 Ollama 위에서 동작하는 qwen3-vl:8b 비전-언어 모델(VLM)을 웹 브라우저에서 "
    "바로 사용할 수 있도록 FastAPI 기반 채팅 웹서비스를 개발하였습니다.")
add_paragraph(doc,
    "텍스트 대화는 물론, 이미지를 업로드하여 시각적 질문 및 분석이 가능하며, 외부에서 "
    "https://atpia.com/qwen3-vl/index.html 로 접속할 수 있습니다.")
add_paragraph(doc,
    "이 서비스는 이전에 개발했던 qwen3:8b 텍스트 전용 웹서비스를 멀티모달로 확장한 버전입니다.")

# ── 2. 개발 배경 ──────────────────────────────────────────────────────────────
add_heading(doc, "2. 개발 배경")
add_divider(doc)
add_paragraph(doc,
    "기존에 운영 중이던 qwen3:8b 채팅 서비스는 텍스트만 처리할 수 있었습니다. "
    "Ollama에 qwen3-vl:8b (Vision-Language) 모델을 설치한 김에, "
    "이미지도 함께 분석할 수 있는 멀티모달 서비스를 별도로 구축하기로 하였습니다.")

add_heading(doc, "2.1 qwen3:8b 대비 주요 차이점", level=2)
add_table(doc,
    ["항목", "qwen3:8b (기존)", "qwen3-vl:8b (이번)"],
    [
        ["입력", "텍스트 전용", "텍스트 + 이미지 (멀티모달)"],
        ["포트", "8000", "8001"],
        ["서브패스", "/qwen/", "/qwen3-vl/"],
        ["이미지 필드", "없음", "messages[].images (base64 배열)"],
        ["타임아웃", "60초", "180초"],
    ]
)

# ── 3. 기술 스택 ──────────────────────────────────────────────────────────────
add_heading(doc, "3. 기술 스택")
add_divider(doc)
add_table(doc,
    ["영역", "사용 기술"],
    [
        ["AI 추론", "Ollama + qwen3-vl:8b (로컬, GPU 활용)"],
        ["백엔드", "Python 3.11, FastAPI, httpx, uvicorn"],
        ["프론트엔드", "HTML5 + Vanilla JS (단일 파일, 프레임워크 없음)"],
        ["통신", "REST API + SSE(Server-Sent Events) 스트리밍"],
        ["이미지 전달", "클라이언트 base64 인코딩 → JSON 페이로드 포함"],
        ["서버", "Ubuntu, Apache 리버스 프록시, systemd"],
        ["GPU", "NVIDIA GTX 1070 8GB"],
    ]
)
add_paragraph(doc,
    "외부 클라우드 API는 일절 사용하지 않고, 모든 추론이 로컬 PC에서 이루어집니다.")

# ── 4. 시스템 구조 ────────────────────────────────────────────────────────────
add_heading(doc, "4. 시스템 구조")
add_divider(doc)
add_code(doc,
    "브라우저 (HTTPS)\n"
    "    │\n"
    "    ▼\n"
    "Apache (atpia.com:443)\n"
    "  ProxyPass /qwen3-vl/ → http://127.0.0.1:8001/\n"
    "    │  flushpackets=on  (SSE 버퍼링 비활성화)\n"
    "    ▼\n"
    "FastAPI / uvicorn (:8001)\n"
    "  POST /api/chat  →  이미지 검증 → Ollama API 호출\n"
    "  GET  /api/health\n"
    "    │\n"
    "    ▼\n"
    "Ollama (:11434)\n"
    "  qwen3-vl:8b 모델 (GPU 추론)"
)

# ── 5. 백엔드 구현 ────────────────────────────────────────────────────────────
add_heading(doc, "5. 백엔드 구현 (main.py)")
add_divider(doc)

add_heading(doc, "5.1 핵심 데이터 모델", level=2)
add_paragraph(doc, "ChatMessage 모델에 images 필드를 추가하여 멀티모달 입력을 지원합니다.")
add_code(doc,
    "class ChatMessage(BaseModel):\n"
    "    role: str\n"
    "    content: str\n"
    "    images: list[str] = Field(default_factory=list)\n"
    "    # base64 인코딩 이미지 배열 (data URI 접두사 포함 가능)"
)

add_heading(doc, "5.2 이미지 처리", level=2)
add_paragraph(doc,
    "Ollama API는 순수 base64만 허용하므로, 브라우저가 전송하는 "
    "data:image/jpeg;base64, 접두사를 strip_data_uri() 함수로 제거합니다.")
add_paragraph(doc,
    "validate_image_size() 함수로 10 MB 초과 이미지를 서버에서 차단하며, "
    "이미지는 서버 디스크에 저장하지 않고 메모리에서만 처리합니다.")

add_heading(doc, "5.3 SSE 스트리밍", level=2)
add_paragraph(doc,
    "Ollama의 스트리밍 응답을 Server-Sent Events로 브라우저에 실시간 전달합니다. "
    "httpx.AsyncClient의 stream() 메서드를 사용하여 비동기로 처리합니다.")

# ── 6. 프론트엔드 구현 ────────────────────────────────────────────────────────
add_heading(doc, "6. 프론트엔드 구현 (static/index.html)")
add_divider(doc)
add_paragraph(doc,
    "단일 HTML 파일로 모든 UI와 로직을 담았습니다. 외부 라이브러리 의존성이 없습니다.")

add_heading(doc, "6.1 이미지 첨부 3가지 방법", level=2)
add_bullet(doc, "📎 버튼: 파일 탐색기에서 선택")
add_bullet(doc, "드래그 & 드롭: 텍스트박스 위에 이미지 파일을 끌어다 놓기")
add_bullet(doc, "클립보드 붙여넣기: 스크린샷 캡처 후 Ctrl+V")

add_heading(doc, "6.2 주요 UI 기능", level=2)
add_bullet(doc, "전송 전 썸네일 미리보기 — 개별 ✕ 버튼으로 제거 가능")
add_bullet(doc, "채팅 버블 내 이미지 표시 — 클릭 시 라이트박스 확대")
add_bullet(doc, "<think> 블록 자동 감지 — 접을 수 있는 '생각 중…' 섹션")
add_bullet(doc, "스트리밍 커서 애니메이션 — 타이핑 효과")
add_bullet(doc, "서브패스 자동 감지 — 리버스 프록시 환경에서도 올바른 API 경로 사용")
add_bullet(doc, "최대 4장 / 10 MB 제한 — 클라이언트 + 서버 양쪽 검증")
add_bullet(doc, "Enter 전송 / Shift+Enter 줄바꿈")

# ── 7. 배포 ───────────────────────────────────────────────────────────────────
add_heading(doc, "7. 배포")
add_divider(doc)

add_heading(doc, "7.1 Apache 리버스 프록시 설정", level=2)
add_paragraph(doc, "atpia.com-le-ssl.conf에 아래 블록을 추가하였습니다.")
add_code(doc,
    "ProxyPreserveHost On\n"
    "ProxyPass        /qwen3-vl/ http://127.0.0.1:8001/ flushpackets=on timeout=185\n"
    "ProxyPassReverse /qwen3-vl/ http://127.0.0.1:8001/\n"
    "LimitRequestBody 15728640"
)
add_bullet(doc, "flushpackets=on: SSE 스트리밍 시 Apache 버퍼링 비활성화")
add_bullet(doc, "timeout=185: uvicorn 180초 타임아웃보다 5초 크게 설정")
add_bullet(doc, "LimitRequestBody 15728640: 15 MB (10 MB + base64 오버헤드 33%)")

add_heading(doc, "7.2 systemd 서비스 + 부팅 자동시작", level=2)
add_paragraph(doc,
    "환경변수는 서비스 파일 내 인라인 대신 별도 .env 파일로 분리하여 관리합니다. "
    "리부팅 후에도 서비스가 자동 시작되도록 아래와 같이 설정하였습니다.")
add_table(doc,
    ["설정 항목", "값", "설명"],
    [
        ["After=", "network-online.target", "네트워크 완전 연결 후 기동"],
        ["Wants=", "ollama.service", "Ollama 먼저 기동 보장"],
        ["EnvironmentFile=", ".env", "환경변수 파일 참조"],
        ["Restart=", "always", "어떤 이유로든 종료 시 자동 재시작"],
        ["TimeoutStartSec=", "60", "부팅 직후 Ollama 로드 시간 여유"],
    ]
)

# ── 8. GitHub 소스코드 공개 ───────────────────────────────────────────────────
add_heading(doc, "8. GitHub 소스코드 공개")
add_divider(doc)
add_paragraph(doc,
    "개발이 완료된 소스코드를 GitHub에 공개하였습니다. "
    "공개 저장소 주소: https://github.com/handuru01/qwen3-vl")

add_heading(doc, "8.1 사전 준비 — .gitignore / .env.example 생성", level=2)
add_paragraph(doc,
    ".env 파일에는 서버 주소 등 환경별 설정이 포함되므로 저장소에서 제외하고, "
    "대신 .env.example 파일을 제공하여 사용자가 참고할 수 있도록 하였습니다.")
add_code(doc,
    "# .gitignore 주요 항목\n"
    ".env          # 실제 환경변수 파일 제외\n"
    "__pycache__/\n"
    "*.pyc\n"
    "venv/"
)
add_code(doc,
    "# .env.example\n"
    "OLLAMA_HOST=http://localhost:11434\n"
    "OLLAMA_MODEL=qwen3-vl:8b\n"
    "REQUEST_TIMEOUT=180\n"
    "MAX_IMAGE_BYTES=10485760\n"
    "APP_PORT=8001"
)

add_heading(doc, "8.2 git 저장소 초기화", level=2)
add_paragraph(doc,
    "처음에는 상위 디렉토리(/mnt/data/handuru/)가 git 루트로 잡혀 파일 경로에 "
    "qwen3-vl/ 접두사가 붙는 문제가 발생하였습니다. "
    "상위 디렉토리의 커밋을 되돌린 후 프로젝트 디렉토리 안에 git 저장소를 새로 초기화하여 해결하였습니다.")
add_code(doc,
    "# 상위 디렉토리의 잘못된 커밋 되돌리기\n"
    "git -C /mnt/data/handuru update-ref -d HEAD\n\n"
    "# 프로젝트 디렉토리에 git 저장소 초기화\n"
    "git init /mnt/data/handuru/qwen3-vl\n"
    "git -C /mnt/data/handuru/qwen3-vl config user.name \"handuru\"\n"
    "git -C /mnt/data/handuru/qwen3-vl config user.email \"hftsys@gmail.com\"\n"
    "git -C /mnt/data/handuru/qwen3-vl branch -m main"
)

add_heading(doc, "8.3 초기 커밋", level=2)
add_paragraph(doc, "11개 파일을 스테이징하고 커밋하였습니다. (.env는 .gitignore에 의해 자동 제외)")
add_code(doc,
    "git add .\n"
    "git commit -m \"Initial commit: Qwen3-VL multimodal chat web service\""
)
add_table(doc,
    ["커밋 포함 파일", "설명"],
    [
        [".env.example",          "환경변수 예시"],
        [".gitignore",            "git 제외 목록"],
        ["CLAUDE.md",             "프로젝트 문서"],
        ["Dockerfile",            "컨테이너 빌드"],
        ["docker-compose.yml",    "Docker Compose 설정"],
        ["apache-qwen3-vl.conf",  "Apache 설정 참고용"],
        ["qwen3-vl.service",      "systemd 서비스 파일"],
        ["main.py",               "FastAPI 백엔드"],
        ["requirements.txt",      "Python 의존성"],
        ["conversation.txt",      "개발 대화 기록"],
        ["static/index.html",     "채팅 UI"],
    ]
)

add_heading(doc, "8.4 GitHub 원격 저장소 연결 및 푸시", level=2)
add_paragraph(doc,
    "https://github.com/new 에서 저장소를 생성한 후 원격 저장소를 연결하고 푸시하였습니다. "
    "GitHub 인증은 Personal Access Token(PAT)을 사용합니다.")
add_code(doc,
    "git remote add origin https://github.com/handuru01/qwen3-vl.git\n"
    "git push -u origin main"
)

add_heading(doc, "8.5 README 작성", level=2)
add_paragraph(doc,
    "GitHub 저장소 메인 화면에 표시될 README.md(영문)와 "
    "README.ko.md(한국어) 두 버전을 작성하였습니다.")
add_bullet(doc, "README.md — 영문 (GitHub 기본 표시)")
add_bullet(doc, "README.ko.md — 한국어")

add_heading(doc, "8.6 docs/ 폴더 구성", level=2)
add_paragraph(doc, "개발보고서 및 문서 파일을 docs/ 폴더에 정리하여 GitHub에 함께 공개하였습니다.")
add_code(doc,
    "docs/\n"
    "├── README.txt                            # 영문 README 텍스트 버전\n"
    "├── qwen3-vl-개발보고서.docx               # 한국어 개발보고서\n"
    "└── qwen3-vl-development report-en-US.docx # 영문 개발보고서"
)

add_heading(doc, "8.7 최종 GitHub 저장소 구성", level=2)
add_code(doc,
    "qwen3-vl/\n"
    "├── README.md                           ← 영문 (GitHub 메인)\n"
    "├── README.ko.md                        ← 한국어\n"
    "├── README.txt\n"
    "├── main.py\n"
    "├── requirements.txt\n"
    "├── .env.example\n"
    "├── .gitignore\n"
    "├── Dockerfile\n"
    "├── docker-compose.yml\n"
    "├── apache-qwen3-vl.conf\n"
    "├── qwen3-vl.service\n"
    "├── make_report.py\n"
    "├── CLAUDE.md\n"
    "├── conversation.txt\n"
    "├── docs/\n"
    "│   ├── README.txt\n"
    "│   ├── qwen3-vl-개발보고서.docx\n"
    "│   └── qwen3-vl-development report- en-US.docx\n"
    "└── static/\n"
    "    └── index.html"
)

# ── 9. 환경변수 ───────────────────────────────────────────────────────────────
add_heading(doc, "9. 환경변수 (.env)")
add_divider(doc)
add_table(doc,
    ["변수명", "기본값", "설명"],
    [
        ["OLLAMA_HOST",     "http://localhost:11434", "Ollama 서버 주소"],
        ["OLLAMA_MODEL",    "qwen3-vl:8b",           "사용 모델명"],
        ["REQUEST_TIMEOUT", "180",                   "요청 타임아웃(초)"],
        ["MAX_IMAGE_BYTES", "10485760",              "업로드 이미지 최대 크기(10MB)"],
        ["APP_PORT",        "8001",                  "uvicorn 기동 포트"],
    ]
)

# ── 10. 프로젝트 구조 ─────────────────────────────────────────────────────────
add_heading(doc, "10. 프로젝트 구조")
add_divider(doc)
add_code(doc,
    "qwen3-vl/\n"
    "├── main.py                 # FastAPI 백엔드 (멀티모달 채팅)\n"
    "├── requirements.txt        # Python 의존성\n"
    "├── .env                    # 환경변수 (서비스 자동시작용, git 제외)\n"
    "├── .env.example            # 환경변수 예시 (GitHub 공개용)\n"
    "├── .gitignore\n"
    "├── Dockerfile\n"
    "├── docker-compose.yml\n"
    "├── apache-qwen3-vl.conf    # Apache 리버스 프록시 설정 참고용\n"
    "├── qwen3-vl.service        # systemd 서비스 파일\n"
    "├── CLAUDE.md               # 프로젝트 문서\n"
    "├── conversation.txt        # 개발 대화 기록\n"
    "└── static/\n"
    "    └── index.html          # 채팅 UI (이미지 업로드 포함)"
)

# ── 11. 동작 확인 ─────────────────────────────────────────────────────────────
add_heading(doc, "11. 동작 확인")
add_divider(doc)
add_paragraph(doc, "헬스체크 API 응답 (curl http://127.0.0.1:8001/api/health):")
add_code(doc,
    '{\n'
    '  "ollama": "ok",\n'
    '  "target_model": "qwen3-vl:8b",\n'
    '  "model_ready": true,\n'
    '  "available_models": ["qwen3-vl:8b", "qwen3:14b", "qwen3:8b"]\n'
    '}'
)
add_paragraph(doc, "브라우저 접속 후 이미지 첨부 및 질의응답 정상 동작 확인 완료.")
add_paragraph(doc, "리부팅 후 서비스 자동 시작 확인 완료.")

# ── 12. 마치며 ────────────────────────────────────────────────────────────────
add_heading(doc, "12. 마치며")
add_divider(doc)
add_paragraph(doc,
    "Ollama 덕분에 GPU 코드 한 줄 없이 로컬에서 멀티모달 AI 서비스를 구동할 수 있었습니다. "
    "FastAPI의 StreamingResponse + SSE 조합이 스트리밍 채팅 구현에 매우 적합하였고, "
    "프론트엔드도 React 같은 프레임워크 없이 Vanilla JS만으로 충분히 구현할 수 있었습니다.")
add_paragraph(doc,
    "외부 API 비용 없이 로컬 GPU 한 장으로 이미지 분석 AI 서비스를 운영할 수 있다는 점이 "
    "가장 큰 장점입니다.")

add_table(doc,
    ["항목", "내용"],
    [
        ["서비스 운영 URL", "https://atpia.com/qwen3-vl/index.html"],
        ["소스코드 공개",   "https://github.com/handuru01/qwen3-vl"],
        ["개발 기간",       "2026년 5월 5일 ~ 2026년 5월 11일"],
        ["개발자",          "김상민"],
    ]
)

# ── 저장 ──────────────────────────────────────────────────────────────────────
output = "/mnt/data/handuru/qwen3-vl/qwen3-vl-개발보고서.docx"
doc.save(output)
print(f"저장 완료: {output}")
