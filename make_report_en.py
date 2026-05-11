"""
Development Report Word Document Generator (English)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(31, 73, 125))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        set_font(run, size=13, bold=True, color=(31, 78, 121))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        set_font(run, size=11, bold=True, color=(68, 68, 68))
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def add_paragraph(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=10)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    rPr.append(shd)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "EAF0FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

    doc.add_paragraph()
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Cover Page ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nQwen3-VL Multimodal AI Chat\nWeb Service Development Report")
set_font(run, size=24, bold=True, color=(31, 73, 125))
p.paragraph_format.space_after = Pt(20)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Building a Web Service for Image + Text Conversation with Local AI")
set_font(run2, size=12, color=(89, 89, 89))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("\n\nDeveloper: Sangmin Kim\nDate: May 11, 2026\n")
set_font(run3, size=11, color=(89, 89, 89))

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("Service URL: https://atpia.com/qwen3-vl/index.html\n"
                   "Source Code: https://github.com/handuru01/qwen3-vl")
set_font(run4, size=10, color=(31, 73, 125))

doc.add_page_break()

# ── 1. Overview ───────────────────────────────────────────────────────────────
add_heading(doc, "1. Overview")
add_divider(doc)
add_paragraph(doc,
    "This project developed a FastAPI-based chat web service that enables direct browser access "
    "to the qwen3-vl:8b vision-language model (VLM) running on a locally installed Ollama instance.")
add_paragraph(doc,
    "In addition to text conversations, users can upload images for visual questions and analysis. "
    "The service is accessible externally at https://atpia.com/qwen3-vl/index.html.")
add_paragraph(doc,
    "This service is a multimodal extension of the previously developed qwen3:8b text-only chat service.")

# ── 2. Background ─────────────────────────────────────────────────────────────
add_heading(doc, "2. Development Background")
add_divider(doc)
add_paragraph(doc,
    "The existing qwen3:8b chat service could only handle text input. "
    "After installing the qwen3-vl:8b (Vision-Language) model on Ollama, "
    "a separate multimodal service capable of analyzing images was built.")

add_heading(doc, "2.1 Key Differences from qwen3:8b", level=2)
add_table(doc,
    ["Item", "qwen3:8b (Previous)", "qwen3-vl:8b (This Project)"],
    [
        ["Input",        "Text only",    "Text + Images (multimodal)"],
        ["Port",         "8000",         "8001"],
        ["Subpath",      "/qwen/",       "/qwen3-vl/"],
        ["Image Field",  "None",         "messages[].images (base64 array)"],
        ["Timeout",      "60 seconds",   "180 seconds"],
    ]
)

# ── 3. Tech Stack ─────────────────────────────────────────────────────────────
add_heading(doc, "3. Tech Stack")
add_divider(doc)
add_table(doc,
    ["Area", "Technology"],
    [
        ["AI Inference",   "Ollama + qwen3-vl:8b (local, GPU-accelerated)"],
        ["Backend",        "Python 3.11, FastAPI, httpx, uvicorn"],
        ["Frontend",       "HTML5 + Vanilla JS (single file, no framework)"],
        ["Communication",  "REST API + SSE (Server-Sent Events) streaming"],
        ["Image Transfer", "Client-side base64 encoding → JSON payload"],
        ["Server",         "Ubuntu, Apache reverse proxy, systemd"],
        ["GPU",            "NVIDIA GTX 1070 8GB"],
    ]
)
add_paragraph(doc,
    "No external cloud APIs are used. All inference runs entirely on the local PC.")

# ── 4. System Architecture ────────────────────────────────────────────────────
add_heading(doc, "4. System Architecture")
add_divider(doc)
add_code(doc,
    "Browser (HTTPS)\n"
    "    │\n"
    "    ▼\n"
    "Apache (atpia.com:443)\n"
    "  ProxyPass /qwen3-vl/ → http://127.0.0.1:8001/\n"
    "    │  flushpackets=on  (disable SSE buffering)\n"
    "    ▼\n"
    "FastAPI / uvicorn (:8001)\n"
    "  POST /api/chat  →  image validation → Ollama API call\n"
    "  GET  /api/health\n"
    "    │\n"
    "    ▼\n"
    "Ollama (:11434)\n"
    "  qwen3-vl:8b model (GPU inference)"
)

# ── 5. Backend Implementation ─────────────────────────────────────────────────
add_heading(doc, "5. Backend Implementation (main.py)")
add_divider(doc)

add_heading(doc, "5.1 Core Data Model", level=2)
add_paragraph(doc,
    "An images field was added to the ChatMessage model to support multimodal input.")
add_code(doc,
    "class ChatMessage(BaseModel):\n"
    "    role: str\n"
    "    content: str\n"
    "    images: list[str] = Field(default_factory=list)\n"
    "    # Array of base64-encoded image strings (data URI prefix allowed)"
)

add_heading(doc, "5.2 Image Processing", level=2)
add_paragraph(doc,
    "Since the Ollama API only accepts pure base64, the strip_data_uri() function removes "
    "the data:image/jpeg;base64, prefix sent by the browser.")
add_paragraph(doc,
    "The validate_image_size() function blocks images exceeding 10 MB on the server side. "
    "Images are processed in memory only and are never saved to disk.")

add_heading(doc, "5.3 SSE Streaming", level=2)
add_paragraph(doc,
    "Ollama's streaming responses are forwarded to the browser in real time using "
    "Server-Sent Events. httpx.AsyncClient's stream() method handles this asynchronously.")

# ── 6. Frontend Implementation ────────────────────────────────────────────────
add_heading(doc, "6. Frontend Implementation (static/index.html)")
add_divider(doc)
add_paragraph(doc,
    "All UI and logic is contained in a single HTML file. No external library dependencies.")

add_heading(doc, "6.1 Three Ways to Attach Images", level=2)
add_bullet(doc, "📎 Button: Select files via the file explorer")
add_bullet(doc, "Drag & Drop: Drop image files onto the text input area")
add_bullet(doc, "Clipboard Paste: Capture a screenshot and press Ctrl+V")

add_heading(doc, "6.2 Key UI Features", level=2)
add_bullet(doc, "Thumbnail preview strip before sending — individual ✕ remove button")
add_bullet(doc, "Image display inside chat bubbles — click to open lightbox")
add_bullet(doc, "Automatic <think> block detection — collapsible 'Thinking...' section")
add_bullet(doc, "Streaming cursor animation — typing effect")
add_bullet(doc, "Subpath auto-detection — correct API path in reverse proxy environments")
add_bullet(doc, "Max 4 images / 10 MB per message — validated on both client and server")
add_bullet(doc, "Enter to send / Shift+Enter for new line")

# ── 7. Deployment ─────────────────────────────────────────────────────────────
add_heading(doc, "7. Deployment")
add_divider(doc)

add_heading(doc, "7.1 Apache Reverse Proxy Configuration", level=2)
add_paragraph(doc, "The following block was added to atpia.com-le-ssl.conf:")
add_code(doc,
    "ProxyPreserveHost On\n"
    "ProxyPass        /qwen3-vl/ http://127.0.0.1:8001/ flushpackets=on timeout=185\n"
    "ProxyPassReverse /qwen3-vl/ http://127.0.0.1:8001/\n"
    "LimitRequestBody 15728640"
)
add_bullet(doc, "flushpackets=on: Disables Apache buffering for SSE streaming")
add_bullet(doc, "timeout=185: Set 5 seconds higher than uvicorn's 180s timeout")
add_bullet(doc, "LimitRequestBody 15728640: 15 MB (10 MB + 33% base64 overhead)")

add_heading(doc, "7.2 systemd Service + Auto-start on Boot", level=2)
add_paragraph(doc,
    "Environment variables are managed in a separate .env file instead of inline in the "
    "service file. The service is configured to start automatically after every reboot.")
add_table(doc,
    ["Setting", "Value", "Description"],
    [
        ["After=",           "network-online.target", "Start after network is fully up"],
        ["Wants=",           "ollama.service",        "Ensure Ollama starts first"],
        ["EnvironmentFile=", ".env",                  "Reference environment file"],
        ["Restart=",         "always",                "Restart on any exit"],
        ["TimeoutStartSec=", "60",                    "Allow time for Ollama to load at boot"],
    ]
)

# ── 8. GitHub Source Code Release ─────────────────────────────────────────────
add_heading(doc, "8. GitHub Source Code Release")
add_divider(doc)
add_paragraph(doc,
    "The completed source code was published to GitHub. "
    "Repository URL: https://github.com/handuru01/qwen3-vl")

add_heading(doc, "8.1 Preparation — .gitignore / .env.example", level=2)
add_paragraph(doc,
    "The .env file contains environment-specific settings and was excluded from the repository. "
    "An .env.example file was provided as a reference template for users.")
add_code(doc,
    "# Key entries in .gitignore\n"
    ".env          # Exclude actual environment file\n"
    "__pycache__/\n"
    "*.pyc\n"
    "venv/"
)
add_code(doc,
    "# .env.example\n"
    "OLLAMA_HOST=http://localhost:11434\n"
    "OLLAMA_MODEL=qwen3-vl:8b\n"
    "REQUEST_TIMEOUT=180\n"
    "MAX_IMAGE_BYTES=10485760\n"
    "APP_PORT=8001"
)

add_heading(doc, "8.2 Git Repository Initialization", level=2)
add_paragraph(doc,
    "Initially, the parent directory (/mnt/data/handuru/) was picked up as the git root, "
    "causing files to be tracked with a qwen3-vl/ prefix. "
    "The parent directory's commit was reverted and a new git repository was initialized "
    "directly inside the project directory.")
add_code(doc,
    "# Revert incorrect commit in parent directory\n"
    "git -C /mnt/data/handuru update-ref -d HEAD\n\n"
    "# Initialize git repository in the correct location\n"
    "git init /mnt/data/handuru/qwen3-vl\n"
    "git -C /mnt/data/handuru/qwen3-vl config user.name \"handuru\"\n"
    "git -C /mnt/data/handuru/qwen3-vl config user.email \"hftsys@gmail.com\"\n"
    "git -C /mnt/data/handuru/qwen3-vl branch -m main"
)

add_heading(doc, "8.3 Initial Commit", level=2)
add_paragraph(doc, "11 files were staged and committed. (.env was automatically excluded by .gitignore)")
add_code(doc,
    "git add .\n"
    "git commit -m \"Initial commit: Qwen3-VL multimodal chat web service\""
)
add_table(doc,
    ["Committed File", "Description"],
    [
        [".env.example",         "Environment variable template"],
        [".gitignore",           "Git exclusion list"],
        ["CLAUDE.md",            "Project documentation"],
        ["Dockerfile",           "Container build"],
        ["docker-compose.yml",   "Docker Compose configuration"],
        ["apache-qwen3-vl.conf", "Apache config reference"],
        ["qwen3-vl.service",     "systemd service file"],
        ["main.py",              "FastAPI backend"],
        ["requirements.txt",     "Python dependencies"],
        ["conversation.txt",     "Development session log"],
        ["static/index.html",    "Chat UI"],
    ]
)

add_heading(doc, "8.4 Connecting Remote Repository and Pushing", level=2)
add_paragraph(doc,
    "A repository was created at https://github.com/new, then the remote was linked and pushed. "
    "GitHub authentication uses a Personal Access Token (PAT).")
add_code(doc,
    "git remote add origin https://github.com/handuru01/qwen3-vl.git\n"
    "git push -u origin main"
)

add_heading(doc, "8.5 README Files", level=2)
add_paragraph(doc,
    "Two versions of the README were written for the GitHub repository main page.")
add_bullet(doc, "README.md — English (displayed by default on GitHub)")
add_bullet(doc, "README.ko.md — Korean")

add_heading(doc, "8.6 docs/ Folder Structure", level=2)
add_paragraph(doc,
    "Development reports and documentation were organized in a docs/ folder "
    "and published to GitHub.")
add_code(doc,
    "docs/\n"
    "├── README.txt                              # English README (text version)\n"
    "├── qwen3-vl-개발보고서.docx                # Korean development report\n"
    "└── qwen3-vl-development report-en-US.docx  # English development report"
)

add_heading(doc, "8.7 Final GitHub Repository Structure", level=2)
add_code(doc,
    "qwen3-vl/\n"
    "├── README.md                           ← English (GitHub main)\n"
    "├── README.ko.md                        ← Korean\n"
    "├── README.txt\n"
    "├── main.py\n"
    "├── requirements.txt\n"
    "├── .env.example\n"
    "├── .gitignore\n"
    "├── Dockerfile\n"
    "├── docker-compose.yml\n"
    "├── apache-qwen3-vl.conf\n"
    "├── qwen3-vl.service\n"
    "├── make_report.py\n"
    "├── make_report_en.py\n"
    "├── CLAUDE.md\n"
    "├── conversation.txt\n"
    "├── docs/\n"
    "│   ├── README.txt\n"
    "│   ├── qwen3-vl-개발보고서.docx\n"
    "│   └── qwen3-vl-development report-en-US.docx\n"
    "└── static/\n"
    "    └── index.html"
)

# ── 9. Environment Variables ──────────────────────────────────────────────────
add_heading(doc, "9. Environment Variables (.env)")
add_divider(doc)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["OLLAMA_HOST",      "http://localhost:11434", "Ollama server address"],
        ["OLLAMA_MODEL",     "qwen3-vl:8b",           "Model name"],
        ["REQUEST_TIMEOUT",  "180",                   "Request timeout (seconds)"],
        ["MAX_IMAGE_BYTES",  "10485760",              "Max upload image size (10MB)"],
        ["APP_PORT",         "8001",                  "uvicorn port"],
    ]
)

# ── 10. Project Structure ─────────────────────────────────────────────────────
add_heading(doc, "10. Project Structure")
add_divider(doc)
add_code(doc,
    "qwen3-vl/\n"
    "├── main.py                 # FastAPI backend (multimodal chat)\n"
    "├── requirements.txt        # Python dependencies\n"
    "├── .env                    # Environment variables (excluded from git)\n"
    "├── .env.example            # Environment variable template (for GitHub)\n"
    "├── .gitignore\n"
    "├── Dockerfile\n"
    "├── docker-compose.yml\n"
    "├── apache-qwen3-vl.conf    # Apache reverse proxy config reference\n"
    "├── qwen3-vl.service        # systemd service file\n"
    "├── CLAUDE.md               # Project documentation\n"
    "├── conversation.txt        # Development session log\n"
    "└── static/\n"
    "    └── index.html          # Chat UI (with image upload)"
)

# ── 11. Verification ──────────────────────────────────────────────────────────
add_heading(doc, "11. Verification")
add_divider(doc)
add_paragraph(doc, "Health check API response (curl http://127.0.0.1:8001/api/health):")
add_code(doc,
    '{\n'
    '  "ollama": "ok",\n'
    '  "target_model": "qwen3-vl:8b",\n'
    '  "model_ready": true,\n'
    '  "available_models": ["qwen3-vl:8b", "qwen3:14b", "qwen3:8b"]\n'
    '}'
)
add_paragraph(doc, "Confirmed normal operation: image attachment and Q&A via browser.")
add_paragraph(doc, "Confirmed automatic service start after system reboot.")

# ── 12. Conclusion ────────────────────────────────────────────────────────────
add_heading(doc, "12. Conclusion")
add_divider(doc)
add_paragraph(doc,
    "Thanks to Ollama, a multimodal AI service was deployed locally without writing a single "
    "line of GPU code. FastAPI's StreamingResponse combined with SSE proved to be an excellent "
    "fit for streaming chat, and the frontend was fully implemented using plain Vanilla JS "
    "without any framework.")
add_paragraph(doc,
    "The greatest advantage is the ability to run an image-analysis AI service on a single "
    "local GPU with zero external API costs.")

add_table(doc,
    ["Item", "Details"],
    [
        ["Service URL",      "https://atpia.com/qwen3-vl/index.html"],
        ["Source Code",      "https://github.com/handuru01/qwen3-vl"],
        ["Development Period", "May 5, 2026 — May 11, 2026"],
        ["Developer",        "Sangmin Kim"],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
output = "/mnt/data/handuru/qwen3-vl/docs/qwen3-vl-development report-en-US.docx"
doc.save(output)
print(f"Saved: {output}")
